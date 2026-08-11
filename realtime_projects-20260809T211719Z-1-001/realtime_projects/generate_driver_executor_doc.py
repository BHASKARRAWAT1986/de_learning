from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

C_HEADER  = "1F497D"
C_DRIVER  = "FFF2CC"
C_EXEC    = "E2EFDA"
C_BOTH    = "DAE3F3"
C_PROBLEM = "FCE4D6"
C_CODE_BG = "F2F2F2"
C_TITLE   = (31, 73, 125)
C_RED     = (192, 0, 0)
C_GREEN   = (0, 112, 0)
C_BLUE    = (68, 114, 196)
C_WHITE   = (255, 255, 255)
C_GREY    = RGBColor(50, 50, 50)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def sf(run, bold=False, italic=False, size=11, color=None, mono=False):
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.name = "Courier New" if mono else "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color) if isinstance(color, tuple) else color

def h(text, level=1, color=C_TITLE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = RGBColor(*color)
        r.font.name = "Calibri"
    return p

def para(*parts):
    p = doc.add_paragraph()
    for part in parts:
        if isinstance(part, str):
            r = p.add_run(part); sf(r)
        else:
            text = part[0]
            bold = part[1] if len(part) > 1 else False
            mono = part[2] if len(part) > 2 else False
            r = p.add_run(text); sf(r, bold=bold, mono=mono)
    return p

def bullet(text, bold_start=None, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + indent*0.3)
    if bold_start and text.startswith(bold_start):
        r1 = p.add_run(bold_start); sf(r1, bold=True)
        rest = text[len(bold_start):]
        if rest:
            r2 = p.add_run(rest); sf(r2, size=10)
    else:
        r = p.add_run(text); sf(r, size=10)
    return p

def code(text, indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  C_CODE_BG)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(8.5); r.font.color.rgb = C_GREY
    return p

def label_box(text, bg):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg)
    pPr.append(shd)
    r = p.add_run(f"  {text}  ")
    sf(r, bold=True, size=10)
    return p

def tbl_header(table, *texts, bg=C_HEADER):
    for j, t in enumerate(texts):
        shade_cell(table.rows[0].cells[j], bg)
        r = table.rows[0].cells[j].paragraphs[0].add_run(t)
        sf(r, bold=True, size=10, color=C_WHITE)
        table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def tbl_cell(table, row, col_idx, text, bg=None, bold=False, mono=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell = table.rows[row].cells[col_idx]
    if bg: shade_cell(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    sf(r, bold=bold, mono=mono, size=9)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Driver vs Executor — Code Analysis")
sf(r, bold=True, size=22, color=C_TITLE)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("bronze-to-silver-insert-py-task  |  Root Cause: High Driver Memory")
sf(r, italic=True, size=12, color=C_BLUE)
doc.add_paragraph()

# Legend
legend = doc.add_table(rows=4, cols=2)
legend.style = "Table Grid"
for i, (bg, label) in enumerate([
    (C_DRIVER,  "🟡  DRIVER  — runs on master node / Spark driver process"),
    (C_EXEC,    "🟢  EXECUTOR — distributed, runs on worker nodes"),
    (C_BOTH,    "🔵  DRIVER (plan) → EXECUTOR (data work)"),
    (C_PROBLEM, "🔴  PROBLEM — directly causes high driver memory"),
]):
    shade_cell(legend.rows[i].cells[0], bg)
    shade_cell(legend.rows[i].cells[1], bg)
    r = legend.rows[i].cells[1].paragraphs[0].add_run(label)
    sf(r, bold=True, size=10)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — CODE WALK-THROUGH
# ══════════════════════════════════════════════════════════════════════════════
h("Part 1 — Full Code Walk-Through: What Runs Where", level=1)

# --- Block 1 ---
h("Block 1 — Imports & Setup", level=2)
label_box("🟡 DRIVER", C_DRIVER)
code("from databricks.sdk.runtime import dbutils\ndbutils.library.restartPython()\nimport os, json, sys, requests\nfrom requests.auth import HTTPBasicAuth\nfrom pyspark.sql.functions import col, udf, from_json ...\nfrom delta.tables import DeltaTable")
para("Everything at the module level runs once on the ", ("DRIVER", True), " at notebook startup. No distribution happens. The driver Python process loads all libraries into its own heap. Executors only get copies when a UDF closure forces it.")

# --- Block 2 ---
h("Block 2 — Config & Schema Load", level=2)
label_box("🟡 DRIVER", C_DRIVER)
code("cfg = JobParameters.get_bronze_to_silver_params()\n_decryption_url      = cfg.decryption_url\n_decryption_user     = cfg.decryption_user\n_decryption_password = cfg.decryption_password\n\n# Dynamic schema load\ntransaction_schema = transaction_schema_module.schema")
para("Config fetched entirely on driver. The three primitive strings are extracted from ", ("cfg", False, True), " specifically so UDFs can be pickled and sent to executors — the full ", ("cfg", False, True), " object references ", ("job_parameters", False, True), " which executors cannot import. The ", ("transaction_schema", False, True), " (deeply nested StructType) lives in driver memory and is referenced by both ", ("from_json()", False, True), " and the return type of the decryption UDF.")

# --- Block 3 ---
h("Block 3 — decrypt_value + udf() registration", level=2)
label_box("🟡 DRIVER (definition)  |  🟢 EXECUTOR (called inside decrypt_transaction_fields)", C_BOTH)
code("def decrypt_value(encrypted_value):\n    response = requests.get(\n        _decryption_url,\n        headers={'x-StringToDecrypt': encrypted_value},\n        auth=HTTPBasicAuth(_decryption_user, _decryption_password),\n        timeout=10\n    )\n    return response.text\n\ndecrypt_udf = udf(decrypt_value, StringType())   # ← registered but NEVER used on a DataFrame column")
para(("Key observation: ", True), "decrypt_udf is defined but never applied to a DataFrame column anywhere in the code. The plain Python function ", ("decrypt_value()", False, True), " IS called directly inside ", ("decrypt_transaction_fields()", False, True), " — meaning HTTP calls happen from within a row-by-row UDF on executors, with zero batching.")

# --- Block 4 ---
h("Block 4 — decrypt_transaction_fields  ← Main Executor Memory Issue", level=2)
label_box("🔴 PROBLEM — Row-by-row Python UDF on full nested struct with HTTP calls per field", C_PROBLEM)
code("def decrypt_transaction_fields(transaction_row):\n    transaction_dict = transaction_row.asDict(recursive=True)  # full deep-copy\n\n    # HTTP call 1 (if GiftCard present):\n    transaction_dict['GiftCard']['AccountNumber'] = decrypt_value(...)\n\n    # HTTP call per item in GiftCards[]:\n    for gift_card in transaction_dict.get('GiftCards', []):\n        gift_card['AccountNumber'] = decrypt_value(...)\n\n    # HTTP call per item in DeletedGiftCards[]:\n    for gift_card in transaction_dict.get('DeletedGiftCards', []):\n        gift_card['AccountNumber'] = decrypt_value(...)\n\n    # HTTP call per tender in Tenders[]:\n    for tender in transaction_dict.get('Tenders', []):\n        tender['AccountNumber'] = decrypt_value(...)\n\n    return transaction_dict\n\ndecrypt_transaction_udf = udf(decrypt_transaction_fields, transaction_schema)  # row-by-row")

doc.add_paragraph()
para(("What happens per row — concrete example with 1 GiftCard + 3 Tenders:", True))

step_tbl = doc.add_table(rows=7, cols=3)
step_tbl.style = "Table Grid"
tbl_header(step_tbl, "Step", "What Happens", "Memory / Time Cost")
step_rows = [
    ("1", "Spark JVM picks one Row from a partition. The entire transaction struct (20–100 KB nested JSON) is pickled: JVM bytes → Python pickle format.", "~20–100 KB per row serialized in executor JVM", C_BOTH),
    ("2", "Python subprocess on executor receives pickle. transaction_row.asDict(recursive=True) creates a FULL deep-copy of the whole nested dict in Python heap.", "Same 20–100 KB now duplicated in Python heap. With 8 cores = 8 simultaneous copies on one executor.", C_EXEC),
    ("3", "decrypt_value() called — HTTP GET to decryption API.\nThis is a BLOCKING call. The executor thread waits for response (up to 10s timeout).\nRepeated for GiftCard + 3 Tenders = 4 sequential HTTP calls.", "4 × network round-trip latency per row.\n4 × HTTP response objects in Python memory.", C_PROBLEM),
    ("4", "Modified dict returned. Python→JVM serialization: dict pickled back to JVM bytes. Spark reconstructs a Row from the returned dict.", "Another 20–100 KB crossing the JVM↔Python boundary.", C_EXEC),
    ("5", "Steps 1–4 repeat for every single row in every partition.\nWith 8 cores per executor, 4 executors = 32 simultaneous Python threads, each holding their own deep-copy + HTTP responses.", "Peak executor Python memory: 32 × ~100 KB = ~3 MB just for in-flight rows (plus Python interpreter base ~50 MB per process).", C_PROBLEM),
    ("6", "If any HTTP call times out or the executor fails, Spark retries the entire task. The retry re-runs ALL HTTP calls for all rows in that partition from scratch.", "Retry multiplies HTTP cost. Driver logs task failures → more driver metadata memory.", C_PROBLEM),
]
for i, (step, what, cost, bg) in enumerate(step_rows, start=1):
    tbl_cell(step_tbl, i, 0, step, bg=bg, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    tbl_cell(step_tbl, i, 1, what, bg=bg)
    tbl_cell(step_tbl, i, 2, cost, bg=bg, bold=(bg==C_PROBLEM))

doc.add_paragraph()

# --- Block 5 ---
h("Block 5 — readStream → select → filter → withColumn", level=2)
label_box("🔵 DRIVER builds plan  |  🟢 EXECUTOR executes per micro-batch", C_BOTH)
code("bronze_stream = spark.readStream.format('delta').table(cfg.bronze_table_full)\n\nsilver_df = bronze_stream\n    .select( col('store_number'), from_json(col('transaction').cast('string'), transaction_schema) )\n    .filter( col('transaction.Header').isNotNull() )\n    .withColumn('transaction_decrypted', decrypt_transaction_udf(col('transaction')))\n    .drop('transaction').withColumnRenamed('transaction_decrypted', 'transaction')")
para("Pure DAG definition — nothing executes yet. The entire plan is stored as a Catalyst logical plan on the driver. ", ("from_json", False, True), " with ", ("transaction_schema", False, True), " correctly parses JSON on executors. The UDF is chained into the plan and will fire on executors per micro-batch.")

# --- Block 6 ---
h("Block 6 — dropDuplicates without watermark", level=2)
label_box("🔴 PROBLEM — Unbounded state store, driver checkpoint thread accumulates work over time", C_PROBLEM)
code("deduped_batch = silver_df.dropDuplicates([\n    'store_number', 'register_number', 'transaction_number', 'transaction_date'\n])\n# ❌ No watermark — Spark must remember EVERY key combination it has EVER seen")
para("In Structured Streaming, ", ("dropDuplicates without watermark", True), " forces Spark to maintain a state store of all unique key combinations since the stream started — growing forever:")
bullet("Day 1 — 100,000 transactions → state = 100K keys (~5 MB)")
bullet("Week 1 — 3.5M transactions → state = 3.5M keys (~175 MB per executor state store)")
bullet("Month 1 — state = GBs — each checkpoint write blocks the driver thread for 30–120 seconds")
bullet("Driver tracks all state store offset metadata → more state = more driver memory for coordination")

doc.add_paragraph()

# --- Block 7 ---
h("Block 7 — write_to_silver()  ← PRIMARY DRIVER MEMORY CULPRIT", level=2)
label_box("🔴 PROBLEM — foreachBatch runs entirely on DRIVER", C_PROBLEM)
code("def write_to_silver(batch_df, batch_id):\n    print(f'Batch {batch_id}: Processing {batch_df.count()} records')  # ← LINE A\n    spark.sql(f'DESCRIBE TABLE {cfg.silver_table_full}')             # ← LINE B\n    delta_table = DeltaTable.forName(spark, cfg.silver_table_full)  # ← LINE C\n    delta_table.alias('target').merge(batch_df.alias('source'), '...').execute()  # ← LINE D")
para(("Critical fact: ", True), "foreachBatch callbacks run entirely on the DRIVER, not on executors. The driver thread calls this function every 2 minutes. Everything inside it happens on the driver.")
doc.add_paragraph()

h("LINE A — batch_df.count()  ← #1 Driver Memory & Performance Killer", level=3, color=C_RED)
code("print(f'Batch {batch_id}: Processing {batch_df.count()} records for Silver')")
para(("Exact sequence of what happens:", True))

count_tbl = doc.add_table(rows=8, cols=2)
count_tbl.style = "Table Grid"
tbl_header(count_tbl, "What happens when batch_df.count() is called inside foreachBatch", "Impact")
count_rows = [
    ("batch_df is a static DataFrame for this micro-batch.\nIt has the FULL DAG attached: Bronze read → from_json → filters → decrypt_transaction_udf.",
     "The complete Catalyst plan (including UDF closure with URLs/credentials + transaction_schema) lives in driver JVM memory.",
     C_DRIVER),
    ("Calling .count() submits a NEW Spark job to the executors.\nThis re-executes the entire plan including decrypt_transaction_udf.",
     "decrypt_value() HTTP calls fire AGAIN for every row in the batch.\nExample: 10,000 rows × 4 fields = 40,000 extra HTTP calls just to count rows.",
     C_PROBLEM),
    ("While executors are running the count job, the driver thread is BLOCKED waiting.",
     "Driver cannot proceed to the actual write.\nThe 2-minute trigger timer is still running — next batch may queue up.",
     C_PROBLEM),
    ("Executors return partial counts. Driver aggregates them into a single Long.",
     "Driver receives the result. Peak memory: BOTH the count job plan AND the subsequent merge plan are in driver JVM simultaneously.",
     C_PROBLEM),
    ("The final result is ONE integer — e.g., 9,847.",
     "The entire re-execution was for a single print() statement.\nZero business value. Pure observational logging.",
     C_PROBLEM),
    ("If count() + merge() together take >2 minutes, the next trigger fires while write_to_silver is still running.",
     "Two batches' query plans accumulate in driver JVM simultaneously.\nDriver heap pressure compounds.",
     C_PROBLEM),
    ("This happens every 2 minutes, indefinitely.",
     "Sustained driver CPU + memory load with no idle time for GC.\nDriver GC pauses grow over time → latency spikes observed as 'high driver usage'.",
     C_PROBLEM),
]
for i, (what, impact, bg) in enumerate(count_rows, start=1):
    tbl_cell(count_tbl, i, 0, what, bg=bg)
    tbl_cell(count_tbl, i, 1, impact, bg=bg, bold=(bg==C_PROBLEM))

doc.add_paragraph()
h("LINE B — DESCRIBE TABLE on every batch", level=3, color=C_RED)
code("spark.sql(f'DESCRIBE TABLE {cfg.silver_table_full}')")
para("Driver makes a network call to Unity Catalog metastore every 2 minutes. After the very first batch creates the table, this check is ALWAYS True. The driver is making 30 unnecessary metastore calls per hour. The DESCRIBE result string is held in driver Python memory until GC collects it.")

doc.add_paragraph()
h("LINE C — DeltaTable.forName() on every batch", level=3, color=C_RED)
code("delta_table = DeltaTable.forName(spark, cfg.silver_table_full)")
para("Every batch, the driver reads the Silver table's Delta transaction log to construct a DeltaTable object. As the table ages and accumulates more writes (each merge = new Delta log entry), the log snapshot the driver reads grows larger. This object is held in driver JVM memory for the duration of the merge.")

doc.add_paragraph()
h("LINE D — merge().execute()  ← #2 Driver Memory Issue", level=3, color=C_RED)
code("delta_table.alias('target').merge(\n    batch_df.alias('source'),\n    'target.transaction_date = source.transaction_date AND\n     target.store_number     = source.store_number     AND\n     target.register_number  = source.register_number  AND\n     target.transaction_number = source.transaction_number'\n).whenNotMatchedInsertAll().execute()")
para("The merge Catalyst plan is compiled and held on the driver JVM. This plan joins:\n• The entire Silver Delta table (potentially 100M+ rows, deeply nested transaction struct with Tenders, GiftCards, LineItems, Header...)\n• The incoming batch_df\n\nFor a complex POS transaction schema with 50–200 nested fields, the Catalyst logical plan + physical plan can be 100–500 MB in driver JVM heap. This plan is held for the entire duration of ", ("execute()", False, True), " — which could be 30–90 seconds. Combined with the count() plan still in memory from LINE A, driver heap pressure peaks here.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — CONCRETE TIMELINE
# ══════════════════════════════════════════════════════════════════════════════
h("Part 2 — Concrete Timeline: One 2-Minute Batch", level=1)
para("Assumption: 10,000 rows per batch, 4 encrypted fields per row, 8 cores × 4 executors = 32 parallel tasks")
doc.add_paragraph()

timeline_tbl = doc.add_table(rows=11, cols=3)
timeline_tbl.style = "Table Grid"
tbl_header(timeline_tbl, "Time", "Activity", "Driver Memory State")

timeline = [
    ("t = 0:00", "Trigger fires. Driver schedules executor tasks to read new Delta files from Bronze table.",
     "Catalyst plan in driver JVM: ~50–200 MB\n(complex schema + UDF closure)", C_DRIVER),
    ("t = 0:05", "Executors process rows:\n• from_json parses nested JSON\n• decrypt_transaction_udf fires per row\n• 32 Python subprocesses each making 4 HTTP calls\n• 32 × 100 KB structs in Python heap simultaneously",
     "Driver holds task metadata for 32 in-flight tasks. Small but accumulating.", C_EXEC),
    ("t = 0:45", "HTTP calls are the bottleneck — each takes ~100ms–2s.\n10,000 rows ÷ 32 parallel = ~313 rows per task.\n313 rows × 4 HTTP calls × avg 300ms = ~375 seconds.\nBatch is already behind schedule.",
     "Driver waiting. Executor tasks accumulate in task scheduler queue on driver.", C_BOTH),
    ("t = 1:30", "Executors finish UDF batch. dropDuplicates state check.\nCheckpoint written to storage (state store).",
     "Driver checkpoint thread writes state store metadata.\nIf state = 5M keys → checkpoint = ~250 MB write to ADLS.", C_PROBLEM),
    ("t = 1:35", "write_to_silver() called on DRIVER.\nbatch_df.count() submitted — FULL DAG RERUNS on executors.\nAnother 10,000 × 4 = 40,000 HTTP calls start.",
     "Driver heap PEAKS:\n• Original batch plan: ~150 MB\n• Count job plan: ~150 MB (same DAG, new job)\n• Waiting for count result\nTotal: ~300–400 MB for plans alone", C_PROBLEM),
    ("t = 3:15", "count() returns: 9,847\nDriver prints: 'Batch 0: Processing 9,847 records'\nAll that compute for this one log line.",
     "Plans partially released but GC hasn't collected yet.\nDriver heap still ~200–300 MB.", C_PROBLEM),
    ("t = 3:16", "spark.sql(DESCRIBE TABLE) — metastore round-trip on driver.\nDeltaTable.forName() — driver reads Silver Delta log.",
     "Delta log snapshot in driver JVM: 10–80 MB depending on table age.", C_DRIVER),
    ("t = 3:18", "merge().execute() — Catalyst compiles the Silver merge plan.\nPlan held in driver JVM: Silver table join + batch_df join on 4 keys.",
     "Merge plan: 100–500 MB in driver JVM.\nCombined with leftover from count job: driver heap = 300–700 MB.", C_PROBLEM),
    ("t = 4:00", "NEXT TRIGGER FIRES while merge is still running.\nwrite_to_silver() is still executing for batch 0.\nBatch 1 queues on driver as pending StreamingQueryProgress.",
     "TWO batches' plans accumulating on driver.\nDriver heap pressure continues to compound.", C_PROBLEM),
    ("Cumulative\n(hours/days)", "dropDuplicates state grows unbounded.\nDelta log grows with each merge.\nBacklogged batches keep queuing.\nDriver GC pauses grow longer.",
     "DRIVER MEMORY: 1–4+ GB and growing.\nGC pause time > 5s → Spark UI shows 'high driver usage'.\nJob may OOM and restart.", C_PROBLEM),
]
for i, (ts, act, mem, bg) in enumerate(timeline, start=1):
    tbl_cell(timeline_tbl, i, 0, ts, bg=bg, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    tbl_cell(timeline_tbl, i, 1, act, bg=bg)
    tbl_cell(timeline_tbl, i, 2, mem, bg=bg, bold=(bg==C_PROBLEM))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 3 — RANKED SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h("Part 3 — Driver Memory Issues Ranked by Impact", level=1)

rank_tbl = doc.add_table(rows=6, cols=4)
rank_tbl.style = "Table Grid"
tbl_header(rank_tbl, "Rank", "Code", "Why it causes HIGH DRIVER MEMORY", "Severity")

ranked = [
    ("1\n(Critical)",
     "batch_df.count()\ninside foreachBatch",
     "foreachBatch runs on DRIVER. count() re-executes full DAG (incl. all UDF HTTP calls) every 2 min. Driver holds two complete query plans simultaneously. Count result = one integer. Total wasted compute: 40,000 HTTP calls per batch.",
     "🔴 CRITICAL\nFix this first"),
    ("2\n(High)",
     "DeltaTable.merge()\n.execute()",
     "Catalyst merge plan for Silver table with deeply nested transaction struct (50–200 fields) compiled + held in driver JVM. Grows as Silver table ages (more stats loaded from Delta log). Plan held 30–90s per batch while executors work.",
     "🔴 HIGH"),
    ("3\n(High)",
     "dropDuplicates([])\nno watermark",
     "Streaming dedup state store grows by N new keys every batch, forever. Driver checkpoint thread spends increasing time per batch coordinating state writes to ADLS. After weeks: GBs of state → driver thread blocked for minutes.",
     "🔴 HIGH"),
    ("4\n(Medium)",
     "spark.sql(DESCRIBE TABLE)\nDeltaTable.forName()\nper batch",
     "Driver calls metastore + reads Delta log every 2 minutes. After first batch, DESCRIBE is always True. 30 unnecessary metastore calls/hour. Delta log snapshot size grows as table ages.",
     "⚠️ MEDIUM"),
    ("5\n(Medium)",
     "decrypt_transaction_udf\n(row-by-row, nested struct)",
     "Primarily causes executor memory pressure. Indirectly affects driver via: (1) slow UDF causes batches to back up → pending batch metadata on driver, (2) task retry metadata when HTTP calls fail, (3) each re-execution (from count()) fires all HTTP calls again.",
     "⚠️ MEDIUM\n(executor primary)"),
]
bg_ranks = [C_PROBLEM, C_PROBLEM, C_PROBLEM, C_DRIVER, C_BOTH]
for i, ((rank, loc, why, sev), bg) in enumerate(zip(ranked, bg_ranks), start=1):
    tbl_cell(rank_tbl, i, 0, rank, bg=bg, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    tbl_cell(rank_tbl, i, 1, loc, bg=bg, bold=True, mono=True)
    tbl_cell(rank_tbl, i, 2, why, bg=bg)
    tbl_cell(rank_tbl, i, 3, sev, bg=bg, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 4 — FIXES
# ══════════════════════════════════════════════════════════════════════════════
h("Part 4 — Fixes with Before/After Code", level=1)

h("Fix 1 — Remove batch_df.count()  [Biggest Win — Fix This First]", level=2, color=C_GREEN)
para(("Before (causes full DAG re-execution every 2 minutes):", True))
code("def write_to_silver(batch_df, batch_id):\n    # ❌ Re-runs the entire DAG including all decrypt HTTP calls just to print one number\n    print(f'Batch {batch_id}: Processing {batch_df.count()} records for Silver')")
para(("After (zero extra compute):", True))
code("def write_to_silver(batch_df, batch_id):\n    # ✅ Remove count() entirely. Get metrics from Spark's own progress tracking:\n    pass\n\n# After stream starts, read progress with no extra compute:\nnum_rows = silver_query.lastProgress['numInputRows']   # ✅ Spark already computed this\nprint(f'Last batch input rows: {num_rows}')")
para(("Impact: ", True), "Halves total DAG executions. Eliminates ~40,000 extra HTTP calls per batch. Reduces peak driver heap by one full query plan (~100–200 MB). Immediately reduces 'high driver usage' metric.")

doc.add_paragraph()
h("Fix 2 — Move table_exists check outside foreachBatch", level=2, color=C_GREEN)
para(("Before (metastore call every 2 minutes):", True))
code("def write_to_silver(batch_df, batch_id):\n    try:\n        spark.sql(f'DESCRIBE TABLE {cfg.silver_table_full}')  # ❌ 30 metastore calls/hour\n        table_exists = True\n    except AnalysisException:\n        table_exists = False")
para(("After (check once at startup):", True))
code("# ✅ Check ONCE before stream starts:\ntry:\n    spark.sql(f'DESCRIBE TABLE {cfg.silver_table_full}')\n    _silver_table_exists = True\nexcept AnalysisException:\n    _silver_table_exists = False\n\ndef write_to_silver(batch_df, batch_id):\n    global _silver_table_exists\n    if not _silver_table_exists:\n        batch_df.write.format('delta').saveAsTable(cfg.silver_table_full)\n        _silver_table_exists = True   # ✅ flip once, never check again\n    else:\n        DeltaTable.forName(spark, cfg.silver_table_full) ...")

doc.add_paragraph()
h("Fix 3 — Add watermark to bound dropDuplicates state", level=2, color=C_GREEN)
para(("Before (unbounded state):", True))
code("deduped_batch = silver_df.dropDuplicates([\n    'store_number', 'register_number', 'transaction_number', 'transaction_date'\n])  # ❌ state grows by N keys every batch, forever")
para(("After (state bounded to 24-hour window):", True))
code("# ✅ Watermark bounds state eviction — Spark drops keys older than 24 hours\nsilver_df_wm = silver_df.withWatermark('SILVER_LAYER_TIMESTAMP', '24 hours')\n\ndeduped_batch = silver_df_wm.dropDuplicates([\n    'store_number', 'register_number', 'transaction_number', 'transaction_date'\n])")
para(("Impact: ", True), "State stays constant at ~24 hours of keys instead of all-time. Example: 500K transactions/day → steady state = 500K keys (~25 MB) instead of growing 500K keys/day. Checkpoint writes stay fast → driver thread unblocked faster each batch.")

doc.add_paragraph()
h("Fix 4 — Replace row-by-row UDF with pandas_udf", level=2, color=C_GREEN)
para(("Before (1 JVM↔Python pickle round-trip per row):", True))
code("# Current: plain Python function wrapped with udf() = row-by-row\ndef decrypt_transaction_fields(transaction_row):\n    transaction_dict = transaction_row.asDict(recursive=True)  # full deep-copy per row\n    # HTTP calls per field...\n    return transaction_dict\ndecrypt_transaction_udf = udf(decrypt_transaction_fields, transaction_schema)")
para(("After (Apache Arrow batch transfer — no per-row pickling):", True))
code("from pyspark.sql.functions import pandas_udf\nimport pandas as pd\n\n# ✅ Apply decryption to individual string columns using pandas_udf:\n# Instead of passing the whole struct, extract the fields, decrypt, reassemble.\n@pandas_udf(StringType())\ndef decrypt_series(s: pd.Series) -> pd.Series:\n    def _decrypt(val):\n        if not val: return val\n        r = requests.get(_decryption_url,\n                         headers={'x-StringToDecrypt': val},\n                         auth=HTTPBasicAuth(_decryption_user, _decryption_password),\n                         timeout=10)\n        return r.text if r.status_code == 200 else val\n    return s.map(_decrypt)\n\n# Apply per encrypted column instead of per whole struct:\nsilver_df = silver_df \\\n    .withColumn('transaction.GiftCard.AccountNumber',\n                decrypt_series(col('transaction.GiftCard.AccountNumber')))")
para(("Impact: ", True), "Arrow columnar transfer: JVM sends entire column as one Arrow buffer instead of N individual pickles. Python receives a pd.Series of N values in one call. Eliminates per-row pickle overhead. Python heap holds one column at a time instead of one full struct per row.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PART 5 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h("Part 5 — Quick Reference Table", level=1)

ref_tbl = doc.add_table(rows=18, cols=3)
ref_tbl.style = "Table Grid"
tbl_header(ref_tbl, "Code Block", "Runs On", "Notes")
ref_tbl.columns[0].width = Inches(2.7)
ref_tbl.columns[1].width = Inches(1.3)
ref_tbl.columns[2].width = Inches(3.0)

ref_rows = [
    ("import statements\ndbutils.library.restartPython()", "DRIVER", C_DRIVER, "Module load, once at startup, driver heap"),
    ("cfg = JobParameters.get_bronze_to_silver_params()", "DRIVER", C_DRIVER, "Config fetch entirely on driver"),
    ("_decryption_url/user/password = cfg.xxx", "DRIVER", C_DRIVER, "Primitive extraction so UDFs can be pickled without importing job_parameters"),
    ("transaction_schema = module.schema", "DRIVER", C_DRIVER, "Nested StructType loaded into driver heap — used by from_json + UDF return type"),
    ("spark.conf.set() / spark.conf.get()", "DRIVER", C_DRIVER, "SparkContext API — driver only"),
    ("def decrypt_value(...)\ndecrypt_udf = udf(decrypt_value, ...)", "DRIVER (def)\nEXECUTOR (call)", C_BOTH, "decrypt_udf NEVER used in DataFrame ops. decrypt_value() called as plain function inside next UDF."),
    ("def decrypt_transaction_fields(...)\ndecrypt_transaction_udf = udf(...)", "DRIVER (def)\nEXECUTOR (call)", C_PROBLEM, "Row-by-row Python UDF. Full struct deep-copied per row. 1–5 HTTP calls per row. Main executor memory consumer."),
    ("spark.readStream.format('delta').table(...)", "DRIVER", C_DRIVER, "DAG definition only — no data read yet"),
    ("bronze_stream.select(from_json(col('transaction').cast('string'), transaction_schema))", "DRIVER (plan)\nEXECUTOR (exec)", C_BOTH, "from_json parses JSON on executors — correct. transaction_schema broadcast to executors."),
    (".filter(col('transaction.Header').isNotNull())", "DRIVER (plan)\nEXECUTOR (exec)", C_EXEC, "Predicate pushdown on executors, may skip Delta files"),
    (".withColumn('transaction_decrypted', decrypt_transaction_udf(col('transaction')))", "EXECUTOR", C_PROBLEM, "Python UDF fires here — row-by-row, blocking HTTP calls per field per row"),
    ("silver_df.dropDuplicates([...]) — no watermark", "DRIVER (plan)\nEXECUTOR (exec)", C_PROBLEM, "Unbounded state store. Grows forever. Driver checkpoint thread increasingly blocked."),
    ("writeStream.foreachBatch(write_to_silver).trigger('2 minutes').start()", "DRIVER", C_DRIVER, "Stream coordinator on driver. Driver thread manages all micro-batch scheduling."),
    ("batch_df.count() inside write_to_silver", "DRIVER (trigger)\nEXECUTOR (compute)", C_PROBLEM, "🔴 #1 DRIVER MEMORY. foreachBatch runs on DRIVER. count() re-executes full DAG per batch. Two plans in driver JVM simultaneously."),
    ("spark.sql(DESCRIBE TABLE) inside write_to_silver", "DRIVER", C_PROBLEM, "Metastore call every 2 min. Unnecessary after first batch. 30 calls/hour wasted."),
    ("DeltaTable.forName(...)", "DRIVER", C_PROBLEM, "Delta log snapshot read into driver JVM every batch. Grows as table ages."),
    ("delta_table.merge(...).execute()", "DRIVER (plan+commit)\nEXECUTOR (exec)", C_PROBLEM, "Catalyst merge plan 100–500 MB in driver JVM. Held until execute() completes. #2 driver memory issue."),
]
for i, (code_text, where, bg, note) in enumerate(ref_rows, start=1):
    tbl_cell(ref_tbl, i, 0, code_text, bg=bg, mono=True)
    tbl_cell(ref_tbl, i, 1, where, bg=bg, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    tbl_cell(ref_tbl, i, 2, note, bg=bg)

doc.add_paragraph()

h("Key Takeaway", level=2)
para(
    ("The root cause of HIGH DRIVER MEMORY is not the UDF", True),
    " (that causes executor memory pressure). The driver memory problem comes from:\n\n"
    "  1.  ",
    ("batch_df.count() inside foreachBatch", True, True),
    " — re-runs full DAG including all HTTP calls, holds two query plans in driver JVM, fires every 2 minutes\n\n"
    "  2.  ",
    ("DeltaTable.merge() on a complex nested schema", True, True),
    " — 100–500 MB Catalyst plan in driver JVM held per batch\n\n"
    "  3.  ",
    ("dropDuplicates without watermark", True, True),
    " — state store grows unboundedly, driver checkpoint thread increasingly stalled\n\n"
    "Fix the ",
    ("count()", False, True),
    " call first — it is one line of code and will give the most immediate and largest reduction in driver memory and CPU."
)

output_path = (
    r"c:\One-Drive\OneDrive - Tredence\Desktop\DSG-PERSONAL\DE LEARNING"
    r"\DATABRICKS\realtime_projects\Driver_vs_Executor_Analysis.docx"
)
doc.save(output_path)
print(f"Saved: {output_path}")
